from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docx.text.font import Font
import os

# 临时文件存储路径（后端目录下的temp文件夹）
TEMP_PATH = os.path.join(os.path.dirname(__file__), "temp")
# 确保temp文件夹存在，不存在则创建
if not os.path.exists(TEMP_PATH):
    os.makedirs(TEMP_PATH)

def generate_meeting_word(meeting_info):
    """
    根据结构化信息生成固定格式的会议记录Word文档
    :param meeting_info: 模型提取的会议信息字典
    :return: 生成的Word文件路径
    """
    # 1. 创建新的Word文档
    doc = Document()

    # 2. 添加标题“会议记录”（居中、加粗）
    title = doc.add_heading('会议记录', 0)  # 0级标题（最大）
    title.alignment = WD_TABLE_ALIGNMENT.CENTER  # 标题居中
    # 设置标题字体（可选，增强美观）
    for run in title.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(22)
        run.font.bold = True

    # 3. 创建“会议基础信息”表格（4行4列，按需求合并单元格）
    base_table = doc.add_table(rows=4, cols=4)
    base_table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格居中
    base_table.autofit = True  # 自动适配内容宽度

    # 设置表格单元格样式：垂直居中、字体统一
    for row in base_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "微软雅黑"
                    run.font.size = Pt(11)

    # 填充基础信息表格（按需求模板填写）
    # 行1：会议主题 + 主持人（合并第3、4列）
    row1_cells = base_table.rows[0].cells
    row1_cells[0].text = "会议主题"
    row1_cells[1].text = meeting_info["meeting_topic"]
    row1_cells[2].merge(row1_cells[3])  # 合并第3、4列
    row1_cells[2].text = "主持人"  # 需求模板中主持人列无信息，留空

    # 行2：会议地点（合并第2-4列）
    row2_cells = base_table.rows[1].cells
    row2_cells[0].text = "会议地点"
    row2_cells[1].merge(row2_cells[3])  # 合并第2-4列
    row2_cells[1].text = meeting_info["meeting_location"]

    # 行3：参会人员（合并第2-4列）
    row3_cells = base_table.rows[2].cells
    row3_cells[0].text = "参会人员"
    row3_cells[1].merge(row3_cells[3])  # 合并第2-4列
    row3_cells[1].text = meeting_info["participants"]

    # 行4：会议时长（合并第2-4列）
    row4_cells = base_table.rows[3].cells
    row4_cells[0].text = "会议时间"
    row4_cells[1].merge(row4_cells[3])  # 合并第2-4列
    row4_cells[1].text = meeting_info["meeting_duration"]

    # 4. 添加“会议内容记录”标题（居中、加粗）
    content_title = doc.add_paragraph("会议内容记录")
    content_title.alignment = WD_TABLE_ALIGNMENT.CENTER
    for run in content_title.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(14)
        run.font.bold = True
    doc.add_paragraph()  # 空行分隔

    # 5. 创建“会议内容”表格（N行4列，按议程数量动态生成）
    agenda_count = len(meeting_info["agenda"])
    content_table = doc.add_table(rows=agenda_count + 1, cols=4)  # 首行为表头
    content_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    content_table.autofit = True

    # 设置内容表格样式（与基础表格一致）
    for row in content_table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.2
                for run in paragraph.runs:
                    run.font.name = "微软雅黑"
                    run.font.size = Pt(10)

    # 填充内容表格表头（合并4列）
    header_cell = content_table.rows[0].cells[0]
    for i in range(1, 4):
        header_cell.merge(content_table.rows[0].cells[i])
    header_cell.text = "会议内容记录"

    # 填充议程内容（按模型提取的agenda数组动态生成）
    for idx, agenda in enumerate(meeting_info["agenda"], 1):
        row_cells = content_table.rows[idx].cells[0]
        # 合并当前行的4列（与需求模板一致，内容占满整行）
        for i in range(1, 4):
            row_cells.merge(content_table.rows[idx].cells[i])
        # 拼接议题内容（含标题、负责人、会前准备、参与人员）
        agenda_text = f"""
议题{idx}：{agenda.get('title', '无')}
负责人：{agenda.get('leader', '无')}
会前准备：{agenda.get('preparation', '无')}
参与人员：{agenda.get('participants', '无')}
        """.strip()  # 去除多余空格
        row_cells.text = agenda_text

    # 6. 添加“会前准备事项”（单独段落，加粗）
    prep_para = doc.add_paragraph()
    prep_run = prep_para.add_run("会前准备事项：")
    prep_run.font.bold = True
    prep_para.add_run(f" {meeting_info['global_preparation']}")

    # 7. 保存Word文件到临时目录
    word_filename = "会议记录.docx"
    word_path = os.path.join(TEMP_PATH, word_filename)
    doc.save(word_path)

    return word_path