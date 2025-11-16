from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.text.font import Font
from docx.oxml.shared import OxmlElement, qn
import os

# 临时文件存储路径（后端目录下的temp文件夹）
TEMP_PATH = os.path.join(os.path.dirname(__file__), "temp")
# 确保temp文件夹存在，不存在则创建
if not os.path.exists(TEMP_PATH):
    os.makedirs(TEMP_PATH)

def set_cell_font(cell, font_name, font_size, is_bold=False):
    """设置单元格字体样式"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = is_bold

def add_cell_border(cell):
    """为单元格添加完整边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # 创建边框元素
    tcBorders = OxmlElement('w:tcBorders')
    
    # 定义完整的边框参数
    border_settings = [
        ('top', 'w:top'),         # 上边框
        ('left', 'w:left'),       # 左边框  
        ('bottom', 'w:bottom'),   # 下边框
        ('right', 'w:right')      # 右边框
    ]
    
    for border_name, border_element in border_settings:
        border = OxmlElement(border_element)
        border.set(qn('w:val'), 'single')      # 实线
        border.set(qn('w:sz'), '8')           # 边框宽度
        border.set(qn('w:space'), '0')        # 边距
        border.set(qn('w:color'), '000000')   # 纯黑色边框
        tcBorders.append(border)
    
    tcPr.append(tcBorders)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def generate_meeting_word(meeting_info):
    """
    根据结构化信息生成与图片一致的会议记录Word文档
    :param meeting_info: 模型提取的会议信息字典
    :return: 生成的Word文件路径
    """
    # 1. 创建新的Word文档
    doc = Document()

    # 2. 添加标题"会议记录"（居中、加粗，黑色黑体）
    title = doc.add_heading('会议记录', 0)
    title.alignment = WD_TABLE_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.name = "黑体"  # 设置为黑体
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # 确保为纯黑色

    # 3. 创建与图片一致的表格（5行4列，结构完全匹配）
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 设置表格单元格样式：垂直居中
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 填充表格内容（严格匹配图片字段位置）
    # 第1行：会议主题 + 主持人
    row1_cells = table.rows[0].cells
    row1_cells[0].text = "会议主题"
    row1_cells[1].text = meeting_info.get("meeting_topic", "无")
    row1_cells[2].text = "主持人"
    row1_cells[3].text = ""  # 若模板中主持人无默认值，留空

    # 第2行：会议地点（合并第2-4列）
    row2_cells = table.rows[1].cells
    row2_cells[0].text = "会议地点"
    row2_cells[1].merge(row2_cells[3])
    row2_cells[1].text = meeting_info.get("meeting_location", "无")

    # 第3行：参会人员（合并第2-4列）
    row3_cells = table.rows[2].cells
    row3_cells[0].text = "参会人员"
    row3_cells[1].merge(row3_cells[3])
    row3_cells[1].text = meeting_info.get("participants", "无")

    # 第4行：会议时长（合并第2-4列）
    row4_cells = table.rows[3].cells
    row4_cells[0].text = "会议时长"
    row4_cells[1].merge(row4_cells[3])
    row4_cells[1].text = meeting_info.get("meeting_duration", "无")

    # 为前4行所有单元格添加边框（会议主题、会议地点、参会人员、会议时长、主持人）
    for row in table.rows[:4]:  # 前4行
        for cell in row.cells:
            add_cell_border(cell)

    # 第5行：会议内容记录（合并第1-4列，添加完整方框）
    row5_cells = table.rows[4].cells
    row5_cells[0].merge(row5_cells[3])
    row5_cells[0].text = "会议内容记录"
    
    # 为会议内容记录单元格添加完整的方框边框
    add_cell_border(row5_cells[0])

    # 4. 设置表格字体样式（标签使用黑体黑色，内容使用微软雅黑）
    # 第1行设置
    set_cell_font(row1_cells[0], "黑体", 11, True)      # 会议主题标签
    set_cell_font(row1_cells[1], "微软雅黑", 11)        # 会议主题内容
    set_cell_font(row1_cells[2], "黑体", 11, True)      # 主持人标签
    set_cell_font(row1_cells[3], "微软雅黑", 11)        # 主持人内容
    
    # 第2行设置
    set_cell_font(row2_cells[0], "黑体", 11, True)      # 会议地点标签
    set_cell_font(row2_cells[1], "微软雅黑", 11)        # 会议地点内容
    
    # 第3行设置
    set_cell_font(row3_cells[0], "黑体", 11, True)      # 参会人员标签
    set_cell_font(row3_cells[1], "微软雅黑", 11)        # 参会人员内容
    
    # 第4行设置
    set_cell_font(row4_cells[0], "黑体", 11, True)      # 会议时长标签
    set_cell_font(row4_cells[1], "微软雅黑", 11)        # 会议时长内容
    
    # 第5行设置（标题使用黑体，内容使用微软雅黑）
    set_cell_font(row5_cells[0], "黑体", 11, True)      # 会议内容记录标签

    # 5. 填充会议内容记录区（议题、负责人等）- 在带边框的单元格内显示
    agenda_count = len(meeting_info.get("agenda", []))
    content_cell = row5_cells[0]
    
    if agenda_count > 0:
        # 在单元格内添加议题内容
        for idx, agenda in enumerate(meeting_info["agenda"]):
            agenda_para = content_cell.add_paragraph()
            agenda_text = f"""议题{idx+1}：{agenda.get('title', '无')}
负责人：{agenda.get('leader', '无')}
会前准备：{agenda.get('preparation', '无')}
参与人员：{agenda.get('participants', '无')}"""
            
            agenda_run = agenda_para.add_run(agenda_text)
            agenda_run.font.name = "微软雅黑"
            agenda_run.font.size = Pt(10)
            
            # 议题之间添加分隔
            if idx < agenda_count - 1:
                content_cell.add_paragraph()
    else:
        # 如果没有议题，显示默认文本
        content_para = content_cell.add_paragraph("无具体议题记录")
        content_para.runs[0].font.name = "微软雅黑"
        content_para.runs[0].font.size = Pt(11)

    # 6. 保存Word文件到临时目录
    word_filename = "会议记录.docx"
    word_path = os.path.join(TEMP_PATH, word_filename)
    doc.save(word_path)

    return word_path